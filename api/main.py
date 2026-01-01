# api/main.py - COMPLETE FIXED CODE
import os
import uuid
from datetime import datetime
from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from qdrant_client import QdrantClient
from qdrant_client.http import models
from openai import OpenAI
import google.generativeai as genai
from motor.motor_asyncio import AsyncIOMotorClient
from bson import ObjectId
import pypdf
import io
# 👇 WORD FILE SUPPORT K LIYE
import docx 
from dotenv import load_dotenv

load_dotenv()

app = FastAPI()

# CORS Setup
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# --- CONFIGURATION ---
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
QDRANT_URL = os.getenv("QDRANT_URL")
QDRANT_API_KEY = os.getenv("QDRANT_API_KEY")
QDRANT_COLLECTION_NAME = "my_documents"
MONGODB_URL = os.getenv("MONGODB_URL")
XAI_API_KEY = os.getenv("XAI_API_KEY")
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")

# --- CLIENTS INIT ---
client_openai = OpenAI(api_key=OPENAI_API_KEY)
client = QdrantClient(url=QDRANT_URL, api_key=QDRANT_API_KEY)
mongo_client = AsyncIOMotorClient(MONGODB_URL)
db = mongo_client.rag_db

if GEMINI_API_KEY:
    genai.configure(api_key=GEMINI_API_KEY)

client_xai = OpenAI(
    api_key=XAI_API_KEY,
    base_url="https://api.x.ai/v1",
)

# --- MODELS ---
class LoginRequest(BaseModel):
    email: str
    password: str

# 👇 YEH MODEL FRONTEND SE MATCH HONA CHAHIYE (422 ERROR FIX)
class QueryRequest(BaseModel):
    query: str              # Frontend must send 'query'
    model_name: str = "gpt"
    user_email: str = None  # Frontend must send 'user_email'
    session_id: str = None

# --- ROUTES ---

@app.get("/")
def home():
    return {"message": "RAG AI Backend is Running!"}

@app.post("/signup")
async def signup(user: LoginRequest):
    existing_user = await db.users.find_one({"email": user.email})
    if existing_user:
        raise HTTPException(status_code=400, detail="User already exists")
    
    new_user = {"email": user.email, "password": user.password} 
    await db.users.insert_one(new_user)
    return {"message": "User created", "email": user.email, "token": "dummy-token"}

@app.post("/login")
async def login(user: LoginRequest):
    db_user = await db.users.find_one({"email": user.email})
    if not db_user or db_user["password"] != user.password:
        raise HTTPException(status_code=401, detail="Invalid credentials")
    return {"message": "Login successful", "email": user.email, "token": "dummy-token"}

@app.post("/index")
async def index_document(file: UploadFile = File(...), email: str = None):
    content = await file.read()
    filename = file.filename.lower()
    text = ""

    # --- 👇 FIXED: FILE TYPE CHECKING (500 ERROR FIX) ---
    try:
        if filename.endswith(".pdf"):
            # Handle PDF
            pdf_reader = pypdf.PdfReader(io.BytesIO(content))
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        
        elif filename.endswith(".docx"):
            # Handle Word Document
            try:
                doc = docx.Document(io.BytesIO(content))
                for para in doc.paragraphs:
                    text += para.text + "\n"
            except Exception as e:
                return {"message": "Error reading Word file. Make sure it is not password protected."}
        
        elif filename.endswith(".txt"):
            # Handle Text File
            text = content.decode("utf-8")
            
        else:
            return {"message": "Unsupported file format. Please upload PDF, DOCX, or TXT."}

    except Exception as e:
        print(f"Error reading file: {e}")
        return {"message": f"Error reading file: {str(e)}"}
    
    if not text.strip():
        return {"message": "File is empty or could not extract text."}

    # Text Chunks
    chunk_size = 1000
    chunks = [text[i:i+chunk_size] for i in range(0, len(text), chunk_size)]
    
    points = []
    for i, chunk in enumerate(chunks):
        # OpenAI Embedding
        response = client_openai.embeddings.create(
            input=chunk,
            model="text-embedding-3-small"
        )
        vector = response.data[0].embedding
        
        # Metadata
        payload = {"text": chunk, "filename": file.filename}
        if email:
            payload["user_email"] = email

        points.append(models.PointStruct(
            id=str(uuid.uuid4()),
            vector=vector,
            payload=payload
        ))
    
    # Qdrant Upload
    try:
        client.get_collection(QDRANT_COLLECTION_NAME)
    except:
        client.create_collection(
            collection_name=QDRANT_COLLECTION_NAME,
            vectors_config=models.VectorParams(size=1536, distance=models.Distance.COSINE)
        )
        
    client.upsert(collection_name=QDRANT_COLLECTION_NAME, points=points)
    return {"message": f"Indexed {len(chunks)} chunks for {file.filename}"}

@app.post("/search")
async def search(request: QueryRequest):
    try:
        print(f"DEBUG: Received search request: {request}") # Log to see what frontend sends

        # 1. Embed Query
        query_vector_response = client_openai.embeddings.create(
            input=request.query,
            model="text-embedding-3-small"
        )
        query_vector = query_vector_response.data[0].embedding

        # 2. Search in Qdrant
        search_filter = None
        if request.user_email:
            search_filter = models.Filter(
                must=[
                    models.FieldCondition(
                        key="user_email",
                        match=models.MatchValue(value=request.user_email)
                    )
                ]
            )

        search_result = client.search(
            collection_name=QDRANT_COLLECTION_NAME,
            query_vector=query_vector,
            query_filter=search_filter,
            limit=3
        )

        # 3. Build Context
        context = ""
        for hit in search_result:
            context += hit.payload.get("text", "") + "\n\n"

        if not context:
            return {"response": "I couldn't find relevant info in your uploaded documents.", "session_id": request.session_id}

        # 4. Generate Answer
        system_prompt = f"Answer based on this context:\n\n{context}"
        answer = "Error generating response."

        if request.model_name == "exai":
            try:
                completion = client_xai.chat.completions.create(
                    model="grok-beta",
                    messages=[
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": request.query}
                    ]
                )
                answer = completion.choices[0].message.content
            except Exception as e:
                answer = f"Grok Error: {str(e)}"

        elif request.model_name == "gemini":
            try:
                model = genai.GenerativeModel('gemini-1.5-flash')
                response = model.generate_content(f"{system_prompt}\nUser: {request.query}")
                answer = response.text
            except Exception as e:
                answer = f"Gemini Error: {str(e)}"

        else: # Default GPT
            try:
                completion = client_openai.chat.completions.create(
                    model="gpt-3.5-turbo",
                    messages=[
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": request.query}
                    ]
                )
                answer = completion.choices[0].message.content
            except Exception as e:
                answer = f"OpenAI Error: {str(e)}"

        # 5. Save History
        if request.user_email:
            if not request.session_id:
                new_session = {
                    "user_email": request.user_email,
                    "title": request.query[:30],
                    "created_at": datetime.now(),
                    "messages": []
                }
                res = await db.sessions.insert_one(new_session)
                request.session_id = str(res.inserted_id)

            await db.sessions.update_one(
                {"_id": ObjectId(request.session_id)},
                {"$push": {"messages": {"user": request.query, "bot": answer}}}
            )

        return {"response": answer, "session_id": request.session_id}

    except Exception as e:
        print(f"CRITICAL ERROR: {str(e)}")
        return {"response": f"Server Error: {str(e)}", "session_id": request.session_id}

@app.get("/history/{email}")
async def get_history(email: str):
    sessions = await db.sessions.find({"user_email": email}).sort("created_at", -1).to_list(100)
    for s in sessions:
        s["_id"] = str(s["_id"])
    return sessions

@app.get("/reset_qdrant")
def reset_qdrant():
    try:
        client.delete_collection(QDRANT_COLLECTION_NAME)
        return {"message": "Success! Purana data delete ho gaya. Ab naya upload karein."}
    except Exception as e:
        return {"message": f"Error: {str(e)}"}

'''
import os
import sys
from io import BytesIO
from typing import List, Optional
import io
from datetime import datetime

# FastAPI Imports
from fastapi import FastAPI, UploadFile, File, HTTPException, Depends
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel

# Path fix taaki Render ko files mil jayein
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

try:
    from models import UserSignup, UserLogin, ChatReq
    from auth import verify_access_token
    from database import db, hash_pass, verify_pass, create_token as create_access_token 
except ImportError:
    from api.models import UserSignup, UserLogin, ChatReq
    from api.auth import verify_access_token
    from api.database import db, hash_pass, verify_pass, create_token as create_access_token 

# AI & RAG Imports
from docx import Document as DocxDocument
from pypdf import PdfReader
from langchain_core.documents import Document as LangchainDocument
from langchain_openai import OpenAIEmbeddings
from langchain_qdrant import QdrantVectorStore
from qdrant_client import QdrantClient
from langchain_text_splitters import RecursiveCharacterTextSplitter
from openai import OpenAI
import google.generativeai as genai
from bson import ObjectId

app = FastAPI()

# Frontend Connect: static folder ko serve karne ke liye
#app.mount("/frontend", StaticFiles(directory="static"), name="static")
origins = [
    "http://localhost:5500",  # Local testing ke liye
    "https://rag-ai-chatbot1-jwob.vercel.app", # Aapka Vercel link
    "https://rag-ai-chatbot1.vercel.app",      # Agar aapka main domain hai
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)
# Health Routes
@app.get("/")
async def root(): return {"message": "AI Chatbot API is Live!"}

@app.get("/health")
async def health(): return {"status": "ok"}

# --- Clients Setup ---
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
QDRANT_URL = os.getenv("QDRANT_URL")
QDRANT_API_KEY = os.getenv("QDRANT_API_KEY")
COLLECTION = os.getenv("QDRANT_COLLECTION_NAME", "AbdulWahab")

openai_client = OpenAI(api_key=OPENAI_API_KEY)
genai.configure(api_key=os.getenv("GEMINI_API_KEY"))
gemini_model = genai.GenerativeModel("gemini-1.5-flash")
embeddings = OpenAIEmbeddings(model="text-embedding-3-small", openai_api_key=OPENAI_API_KEY)

# --- Pydantic Models ---
class UserAuth(BaseModel):
    email: str
    password: str

class ChatReq(BaseModel):
    query: str
    model_name: str = "gpt"
    session_id: Optional[str] = None
    user_email: str

# --- 1. Auth Endpoints ---
@app.post("/signup")
async def signup(user: UserAuth):
    if await db.users.find_one({"email": user.email}):
        raise HTTPException(400, "Email already exists")
    await db.users.insert_one({"email": user.email, "password": hash_pass(user.password)})
    return {"status": "success"}

@app.post("/login")
async def login(user: UserAuth):
    db_user = await db.users.find_one({"email": user.email})
    if not db_user or not verify_pass(user.password, db_user["password"]):
        raise HTTPException(400, "Invalid credentials")
    return {"token": create_access_token({"sub": user.email}), "email": user.email}

# --- 2. Indexing Endpoint ---
@app.post("/index")
async def index_doc(file: UploadFile = File(...)):
    try:
        data = await file.read()
        text = ""
        filename = file.filename.lower()

        # --- YAHAN NAYA LOGIC ADD HOGA ---
        if filename.endswith(".pdf"):
            pdf = PdfReader(BytesIO(data))
            for page in pdf.pages: 
                text += page.extract_text() or ""
        
        elif filename.endswith(".docx"):
            doc = DocxDocument(io.BytesIO(data))
            for para in doc.paragraphs:
                text += para.text + "\n"
            
        elif filename.endswith(".txt"):
            text = data.decode("utf-8")

        
        else:
            raise HTTPException(400, "Sirf PDF, DOCX, aur TXT support hain.")
        # --------------------------------

        if not text.strip():
            raise HTTPException(400, "File khali hai.")

        # --- YAHAN SE AAPKA QDRANT WALA PURANA LOGIC SHURU HOGA ---
        splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=150)
        chunks = splitter.split_documents([LangchainDocument(page_content=text, metadata={"source": file.filename})])

        QdrantVectorStore.from_documents(
            documents=chunks, embedding=embeddings,
            url=QDRANT_URL, api_key=QDRANT_API_KEY, collection_name=COLLECTION
        )
        return {"status": "success", "filename": file.filename}
    
    except Exception as e:
        raise HTTPException(500, str(e))

# --- 3. Chat & History Endpoint ---
@app.post("/search")
async def chat(req: ChatReq):
    try:
        # RAG Retrieval
        client = QdrantClient(url=QDRANT_URL, api_key=QDRANT_API_KEY)
        vectorstore = QdrantVectorStore(client=client, collection_name=COLLECTION, embedding=embeddings)
        docs = vectorstore.similarity_search(req.query, k=3)
        context = "\n".join([d.page_content for d in docs])
        
        # Generation
        prompt = f"Context: {context}\n\nQuestion: {req.query}"
        if req.model_name == "gemini":
            res = gemini_model.generate_content(prompt)
            answer = res.texts
        elif req.model_name == "exai":
            from openai import OpenAI as XAIClient
            x_client = XAIClient(
                api_key=os.getenv("XAI_API_KEY"),
                base_url="https://api.x.ai/v1",
            )
            res = x_client.chat.completions.create(
                model="grok-beta",
                messages=[{"role": "user", "content": prompt}]
            )    
            answer = res.choices[0].message.content
        else:
            res = openai_client.chat.completions.create(model="gpt-3.5-turbo", messages=[{"role": "user", "content": prompt}])
            answer = res.choices[0].message.content

        # Save to MongoDB History
        chat_msg = {"user": req.query, "bot": answer, "time": datetime.utcnow()}
        if req.session_id:
            await db.sessions.update_one({"_id": ObjectId(req.session_id)}, {"$push": {"messages": chat_msg}})
            session_id = req.session_id
        else:
            new_s = await db.sessions.insert_one({
                "user_email": req.user_email,
                "title": req.query[:30] + "...",
                "messages": [chat_msg]
            })
            session_id = str(new_s.inserted_id)

        return {"response": answer, "session_id": session_id, "sources": list(set([d.metadata['source'] for d in docs]))}
    except Exception as e:
        raise HTTPException(500, str(e))

@app.get("/history/{email}")
async def get_history(email: str):
    sessions = await db.sessions.find({"user_email": email}).to_list(100)
    for s in sessions: s["_id"] = str(s["_id"])
    return sessions

'''





